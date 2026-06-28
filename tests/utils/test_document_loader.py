import os
from collections.abc import Iterator
from unittest.mock import MagicMock, patch

import pytest

from app.utils.document_loader import get_loader, clean_text, process_documents
from langchain_community.document_loaders import (
    TextLoader,
    UnstructuredMarkdownLoader,
)
from langchain_core.documents import Document


def test_clean_text():
    text = "Hello\x00World"
    cleaned = clean_text(text)
    assert "\x00" not in cleaned
    assert cleaned == "HelloWorld"


def test_get_loader_text(tmp_path):
    # Create a temporary text file.
    file_path = tmp_path / "test.txt"
    file_path.write_text("Sample text")
    loader, known_type, file_ext = get_loader("test.txt", "text/plain", str(file_path))
    assert known_type is True
    assert file_ext == "txt"
    data = loader.load()
    # Check that data is loaded.
    assert data is not None


def test_process_documents():
    docs = [
        Document(
            page_content="Page 1 content", metadata={"source": "dummy.txt", "page": 1}
        ),
        Document(
            page_content="Page 2 content", metadata={"source": "dummy.txt", "page": 2}
        ),
    ]
    processed = process_documents(docs)
    assert "dummy.txt" in processed
    assert "# PAGE 1" in processed
    assert "# PAGE 2" in processed


def test_safe_pdf_loader_class():
    """Test that SafePyPDFLoader class can be instantiated"""
    from app.utils.document_loader import SafePyPDFLoader

    loader = SafePyPDFLoader("dummy.pdf")
    assert loader.filepath == "dummy.pdf"
    assert loader._temp_filepath is None


def test_get_loader_text_lazy_load(tmp_path):
    """Test that lazy_load returns an iterator yielding documents."""
    file_path = tmp_path / "test.txt"
    file_path.write_text("Sample text")
    loader, known_type, file_ext = get_loader("test.txt", "text/plain", str(file_path))
    assert known_type is True
    assert file_ext == "txt"
    data = list(loader.lazy_load())
    assert len(data) > 0
    assert hasattr(data[0], "page_content")


def test_get_loader_pdf(tmp_path):
    """Test get_loader returns SafePyPDFLoader for PDF files"""
    # Create a dummy PDF file path (doesn't need to be real for this test)
    file_path = tmp_path / "test.pdf"
    file_path.write_text("dummy content")  # Not a real PDF, but that's OK for this test

    loader, known_type, file_ext = get_loader(
        "test.pdf", "application/pdf", str(file_path)
    )

    # Check that we get our SafePyPDFLoader
    from app.utils.document_loader import SafePyPDFLoader

    assert isinstance(loader, SafePyPDFLoader)
    assert known_type is True
    assert file_ext == "pdf"


def test_safe_pdf_loader_lazy_load():
    """Test that SafePyPDFLoader.lazy_load() returns an Iterator."""
    from app.utils.document_loader import SafePyPDFLoader

    loader = SafePyPDFLoader("dummy.pdf")
    assert hasattr(loader, "lazy_load")
    result = loader.lazy_load()
    assert isinstance(result, Iterator)


def _make_mistral_module(pages=None, raise_exc=None):
    """Build a fake ``mistralai`` module whose ``Mistral().ocr.process`` is stubbed.

    Returns ``(fake_module, client)`` so callers can both inject the module via
    ``sys.modules`` and make assertions on the OCR client.
    """
    fake_module = MagicMock()
    client = MagicMock()
    if raise_exc is not None:
        client.ocr.process.side_effect = raise_exc
    else:
        response = MagicMock()
        response.pages = pages
        client.ocr.process.return_value = response
    fake_module.Mistral.return_value = client
    return fake_module, client


def test_safe_pdf_loader_ocr_maps_pages():
    """load() maps Mistral OCR pages to Documents with 1-based sequential page metadata."""
    from app.utils import document_loader
    from app.utils.document_loader import SafePyPDFLoader

    # Mistral indexes pages 0-based; we renumber to 1-based in response order.
    page0 = MagicMock(index=0, markdown="page zero text")
    page1 = MagicMock(index=1, markdown="page one text")
    fake_module, client = _make_mistral_module(pages=[page0, page1])

    loader = SafePyPDFLoader("dummy.pdf")
    with patch.dict("sys.modules", {"mistralai": fake_module}), patch.object(
        document_loader, "MISTRAL_API_KEY", "test-key"
    ), patch.object(SafePyPDFLoader, "_encode_pdf_b64", return_value="b64data"):
        result = loader.load()

    assert [d.page_content for d in result] == ["page zero text", "page one text"]
    assert result[0].metadata == {"source": "dummy.pdf", "page": 1}
    assert result[1].metadata == {"source": "dummy.pdf", "page": 2}
    client.ocr.process.assert_called_once()


def test_safe_pdf_loader_lazy_load_yields_ocr_pages():
    """lazy_load() yields the same Documents that load() produces."""
    from app.utils.document_loader import SafePyPDFLoader

    docs = [Document(page_content="p1"), Document(page_content="p2")]
    loader = SafePyPDFLoader("dummy.pdf")
    with patch.object(SafePyPDFLoader, "load", return_value=docs):
        result = list(loader.lazy_load())

    assert result == docs


def test_safe_pdf_loader_empty_pages_returns_placeholder():
    """An OCR response with no pages yields a single empty placeholder Document."""
    from app.utils import document_loader
    from app.utils.document_loader import SafePyPDFLoader

    fake_module, _ = _make_mistral_module(pages=[])

    loader = SafePyPDFLoader("dummy.pdf")
    with patch.dict("sys.modules", {"mistralai": fake_module}), patch.object(
        document_loader, "MISTRAL_API_KEY", "test-key"
    ), patch.object(SafePyPDFLoader, "_encode_pdf_b64", return_value="b64data"):
        result = loader.load()

    assert len(result) == 1
    assert result[0].page_content == ""
    assert result[0].metadata == {"source": "dummy.pdf", "page": 1}


def test_safe_pdf_loader_ocr_error_propagates():
    """An error raised by the Mistral OCR API call propagates to the caller."""
    from app.utils import document_loader
    from app.utils.document_loader import SafePyPDFLoader

    fake_module, _ = _make_mistral_module(raise_exc=RuntimeError("OCR boom"))

    loader = SafePyPDFLoader("dummy.pdf")
    with patch.dict("sys.modules", {"mistralai": fake_module}), patch.object(
        document_loader, "MISTRAL_API_KEY", "test-key"
    ), patch.object(SafePyPDFLoader, "_encode_pdf_b64", return_value="b64data"):
        with pytest.raises(RuntimeError, match="OCR boom"):
            loader.load()


def test_safe_pdf_loader_requires_api_key():
    """load() raises a clear error when MISTRAL_API_KEY is not configured."""
    from app.utils import document_loader
    from app.utils.document_loader import SafePyPDFLoader

    fake_module, _ = _make_mistral_module(pages=[])

    loader = SafePyPDFLoader("dummy.pdf")
    with patch.dict("sys.modules", {"mistralai": fake_module}), patch.object(
        document_loader, "MISTRAL_API_KEY", ""
    ):
        with pytest.raises(RuntimeError, match="MISTRAL_API_KEY"):
            loader.load()


MARKDOWN_SAMPLE = (
    "# Heading\n\n"
    "**bold** and *italic* text with a [link](https://example.com).\n\n"
    "- item 1\n"
    "- item 2\n\n"
    "> a blockquote\n"
)


def test_get_loader_markdown_embed_uses_unstructured(tmp_path):
    """Default (embedding) path must keep UnstructuredMarkdownLoader for .md."""
    file_path = tmp_path / "notes.md"
    file_path.write_text(MARKDOWN_SAMPLE, encoding="utf-8")

    loader, known_type, file_ext = get_loader(
        "notes.md", "text/markdown", str(file_path)
    )

    assert isinstance(loader, UnstructuredMarkdownLoader)
    assert known_type is True
    assert file_ext == "md"


@pytest.mark.parametrize(
    "content_type",
    [
        "text/markdown",
        "text/x-markdown",
        "application/markdown",
        "application/x-markdown",
    ],
)
def test_get_loader_markdown_raw_text_uses_text_loader(tmp_path, content_type):
    """/text path (raw_text=True) must load .md verbatim so formatting survives."""
    file_path = tmp_path / "notes.md"
    file_path.write_text(MARKDOWN_SAMPLE, encoding="utf-8")

    loader, known_type, file_ext = get_loader(
        "notes.md", content_type, str(file_path), raw_text=True
    )

    assert isinstance(loader, TextLoader)
    assert known_type is True
    assert file_ext == "md"

    docs = loader.load()
    assert len(docs) == 1
    assert docs[0].page_content == MARKDOWN_SAMPLE


def test_get_loader_markdown_raw_text_by_extension_only(tmp_path):
    """Extension-based detection must still kick in when content type is generic."""
    file_path = tmp_path / "README.md"
    file_path.write_text(MARKDOWN_SAMPLE, encoding="utf-8")

    loader, _, _ = get_loader(
        "README.md", "application/octet-stream", str(file_path), raw_text=True
    )

    assert isinstance(loader, TextLoader)


def test_get_loader_raw_text_leaves_pdf_alone(tmp_path):
    """raw_text must not disturb binary formats — PDF still uses the PDF loader."""
    from app.utils.document_loader import SafePyPDFLoader

    file_path = tmp_path / "doc.pdf"
    file_path.write_text("not a real pdf")

    loader, _, file_ext = get_loader(
        "doc.pdf", "application/pdf", str(file_path), raw_text=True
    )

    assert isinstance(loader, SafePyPDFLoader)
    assert file_ext == "pdf"


@pytest.mark.parametrize(
    "filename, expected_loader_name",
    [
        ("doc.pdf", "SafePyPDFLoader"),
        # .docx on the /text path (raw_text=True) routes to pandoc for tracked
        # changes/comments; the markdown Content-Type must not override that.
        ("report.docx", "PandocDocxLoader"),
        ("book.epub", "UnstructuredEPubLoader"),
        ("data.xlsx", "UnstructuredExcelLoader"),
        ("slides.pptx", "UnstructuredPowerPointLoader"),
    ],
)
def test_get_loader_raw_text_respects_binary_extensions_over_markdown_mime(
    tmp_path, filename, expected_loader_name
):
    """A markdown Content-Type must not override a known binary extension.

    Some clients send conflicting multipart content types. For an upload named
    `doc.pdf` with Content-Type `text/markdown`, the PDF loader still has to
    win — otherwise a binary file is read as UTF-8 text.
    """
    file_path = tmp_path / filename
    file_path.write_text("placeholder binary content")

    loader, _, _ = get_loader(
        filename, "text/markdown", str(file_path), raw_text=True
    )

    assert type(loader).__name__ == expected_loader_name


# ---------------------------------------------------------------------------
# DOCX via pandoc (/text endpoint)
# ---------------------------------------------------------------------------


def test_get_loader_docx_embed_uses_docx2txt(tmp_path):
    """Embedding path (raw_text=False) keeps Docx2txtLoader for .docx."""
    from langchain_community.document_loaders import Docx2txtLoader

    file_path = tmp_path / "report.docx"
    file_path.write_bytes(b"placeholder")

    loader, known_type, file_ext = get_loader(
        "report.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        str(file_path),
    )
    assert isinstance(loader, Docx2txtLoader)
    assert known_type is True
    assert file_ext == "docx"


def test_get_loader_docx_raw_text_uses_pandoc(tmp_path):
    """/text path (raw_text=True) routes .docx to PandocDocxLoader."""
    from app.utils.document_loader import PandocDocxLoader

    file_path = tmp_path / "report.docx"
    file_path.write_bytes(b"placeholder")

    loader, _, file_ext = get_loader(
        "report.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        str(file_path),
        raw_text=True,
    )
    assert isinstance(loader, PandocDocxLoader)
    assert file_ext == "docx"


def test_get_loader_legacy_doc_is_rejected(tmp_path):
    """Legacy binary .doc can't be loaded by any backend, so it's rejected clearly."""
    file_path = tmp_path / "legacy.doc"
    file_path.write_bytes(b"placeholder")

    with pytest.raises(ValueError, match="Legacy .doc files are not supported"):
        get_loader("legacy.doc", "application/msword", str(file_path))


def test_get_loader_legacy_doc_rejected_even_with_markdown_mime(tmp_path):
    """A markdown Content-Type must not route .doc into the markdown loader either."""
    file_path = tmp_path / "legacy.doc"
    file_path.write_bytes(b"placeholder")

    with pytest.raises(ValueError, match="Legacy .doc files are not supported"):
        get_loader("legacy.doc", "text/markdown", str(file_path), raw_text=True)


def test_pandoc_docx_loader_passes_track_changes_and_builds_document(tmp_path):
    """PandocDocxLoader converts via pypandoc with --track-changes and wraps output."""
    from app.utils.document_loader import PandocDocxLoader

    file_path = tmp_path / "doc.docx"
    file_path.write_bytes(b"placeholder")

    fake_pypandoc = MagicMock()
    fake_pypandoc.convert_file.return_value = "converted [text]{.insertion} markdown"

    loader = PandocDocxLoader(str(file_path), track_changes="all")
    with patch.dict("sys.modules", {"pypandoc": fake_pypandoc}):
        docs = loader.load()

    assert len(docs) == 1
    assert docs[0].page_content == "converted [text]{.insertion} markdown"
    assert docs[0].metadata["source"] == str(file_path)

    _, kwargs = fake_pypandoc.convert_file.call_args
    assert kwargs["to"] == "markdown"
    assert kwargs["format"] == "docx"
    assert "--track-changes=all" in kwargs["extra_args"]


def test_pandoc_docx_loader_lazy_load_yields_from_load(tmp_path):
    from app.utils.document_loader import PandocDocxLoader

    loader = PandocDocxLoader(str(tmp_path / "doc.docx"))
    docs = [Document(page_content="x")]
    with patch.object(PandocDocxLoader, "load", return_value=docs):
        result = list(loader.lazy_load())
    assert result == docs


def _make_docx_with_header_footer(path, header, footer, body="Body paragraph."):
    from docx import Document as Docx

    doc = Docx()
    section = doc.sections[0]
    section.header.paragraphs[0].text = header
    section.footer.paragraphs[0].text = footer
    doc.add_paragraph(body)
    doc.save(path)


def test_extract_docx_headers_footers(tmp_path):
    from app.utils.document_loader import _extract_docx_headers_footers

    p = tmp_path / "doc.docx"
    _make_docx_with_header_footer(
        str(p), "PRIVILEGED & CONFIDENTIAL", "Matter No. 12345"
    )

    block = _extract_docx_headers_footers(str(p))
    assert "[Header] PRIVILEGED & CONFIDENTIAL" in block
    assert "[Footer] Matter No. 12345" in block


def test_extract_docx_headers_footers_handles_bad_file(tmp_path):
    """Non-docx input must not raise; returns an empty string so body load survives."""
    from app.utils.document_loader import _extract_docx_headers_footers

    p = tmp_path / "bad.docx"
    p.write_bytes(b"not a real docx")
    assert _extract_docx_headers_footers(str(p)) == ""


def test_pandoc_docx_loader_prepends_headers_footers(tmp_path):
    from app.utils.document_loader import PandocDocxLoader

    p = tmp_path / "doc.docx"
    _make_docx_with_header_footer(
        str(p), "PRIVILEGED & CONFIDENTIAL", "Matter No. 12345"
    )

    fake_pypandoc = MagicMock()
    fake_pypandoc.convert_file.return_value = "Body markdown from pandoc."

    loader = PandocDocxLoader(str(p), include_headers_footers=True)
    with patch.dict("sys.modules", {"pypandoc": fake_pypandoc}):
        docs = loader.load()

    text = docs[0].page_content
    assert "PRIVILEGED & CONFIDENTIAL" in text
    assert "Matter No. 12345" in text
    assert "Body markdown from pandoc." in text
    # Header/footer block precedes the pandoc body
    assert text.index("PRIVILEGED") < text.index("Body markdown")


def test_pandoc_docx_loader_headers_footers_disabled(tmp_path):
    from app.utils.document_loader import PandocDocxLoader

    p = tmp_path / "doc.docx"
    _make_docx_with_header_footer(str(p), "CONFIDENTIAL", "Matter No. 999")

    fake_pypandoc = MagicMock()
    fake_pypandoc.convert_file.return_value = "Body markdown only."

    loader = PandocDocxLoader(str(p), include_headers_footers=False)
    with patch.dict("sys.modules", {"pypandoc": fake_pypandoc}):
        docs = loader.load()

    assert docs[0].page_content == "Body markdown only."
    assert "CONFIDENTIAL" not in docs[0].page_content


# ---------------------------------------------------------------------------
# RTF (.rtf)
# ---------------------------------------------------------------------------


def test_get_loader_rtf(tmp_path):
    from langchain_community.document_loaders import UnstructuredRTFLoader

    file_path = tmp_path / "brief.rtf"
    file_path.write_text(r"{\rtf1 hello}", encoding="utf-8")

    loader, known_type, file_ext = get_loader(
        "brief.rtf", "application/rtf", str(file_path)
    )
    assert isinstance(loader, UnstructuredRTFLoader)
    assert known_type is True
    assert file_ext == "rtf"


def test_get_loader_rtf_not_hijacked_by_markdown_mime(tmp_path):
    """A markdown Content-Type must not route .rtf into the markdown loader."""
    from langchain_community.document_loaders import UnstructuredRTFLoader

    file_path = tmp_path / "brief.rtf"
    file_path.write_text(r"{\rtf1 hello}", encoding="utf-8")

    loader, _, _ = get_loader("brief.rtf", "text/markdown", str(file_path))
    assert isinstance(loader, UnstructuredRTFLoader)


# ---------------------------------------------------------------------------
# Email (.eml / .msg)
# ---------------------------------------------------------------------------

EML_SAMPLE = (
    "From: Alice <alice@example.com>\n"
    "To: Bob <bob@example.com>\n"
    "Subject: Quarterly numbers\n"
    "Date: Tue, 1 Apr 2025 10:00:00 -0000\n"
    'Content-Type: text/plain; charset="utf-8"\n'
    "\n"
    "Hi Bob,\n"
    "Here are the Q1 results. Revenue up 12%.\n"
    "Thanks,\n"
    "Alice\n"
)


def test_get_loader_eml(tmp_path):
    from app.utils.document_loader import EmailLoader

    file_path = tmp_path / "mail.eml"
    file_path.write_text(EML_SAMPLE, encoding="utf-8")

    loader, known_type, file_ext = get_loader(
        "mail.eml", "message/rfc822", str(file_path)
    )
    assert isinstance(loader, EmailLoader)
    assert known_type is True
    assert file_ext == "eml"


def test_get_loader_eml_not_hijacked_by_markdown_mime(tmp_path):
    """A markdown Content-Type must not route .eml into the markdown loader."""
    from app.utils.document_loader import EmailLoader

    file_path = tmp_path / "mail.eml"
    file_path.write_text(EML_SAMPLE, encoding="utf-8")

    loader, _, _ = get_loader("mail.eml", "text/markdown", str(file_path))
    assert isinstance(loader, EmailLoader)


def test_email_loader_prepends_headers(tmp_path):
    from app.utils.document_loader import EmailLoader

    file_path = tmp_path / "mail.eml"
    file_path.write_text(EML_SAMPLE, encoding="utf-8")

    docs = EmailLoader(str(file_path), include_headers=True).load()
    assert len(docs) == 1
    text = docs[0].page_content
    assert "From: Alice <alice@example.com>" in text
    assert "To: Bob <bob@example.com>" in text
    assert "Subject: Quarterly numbers" in text
    assert "Q1 results" in text
    # Header block precedes the body
    assert text.index("Subject:") < text.index("Q1 results")


def test_email_loader_body_only_when_headers_disabled(tmp_path):
    from app.utils.document_loader import EmailLoader

    file_path = tmp_path / "mail.eml"
    file_path.write_text(EML_SAMPLE, encoding="utf-8")

    docs = EmailLoader(str(file_path), include_headers=False).load()
    text = docs[0].page_content
    assert "From:" not in text
    assert "Subject:" not in text
    assert "Q1 results" in text


def test_get_loader_msg(tmp_path):
    from app.utils.document_loader import OutlookMsgLoader

    file_path = tmp_path / "mail.msg"
    file_path.write_bytes(b"placeholder")

    loader, known_type, file_ext = get_loader(
        "mail.msg", "application/vnd.ms-outlook", str(file_path)
    )
    assert isinstance(loader, OutlookMsgLoader)
    assert file_ext == "msg"


def test_outlook_msg_loader_builds_headers_from_transport_headers(tmp_path):
    """Headers come from transport headers (To/Cc preserved, Bcc absent by nature)."""
    from app.utils.document_loader import OutlookMsgLoader

    file_path = tmp_path / "mail.msg"
    file_path.write_bytes(b"placeholder")

    fake_msg = MagicMock()
    fake_msg.body = "Hi Bob,\nHere are the Q1 results."
    fake_msg.message_headers = {
        "From": "Alice <alice@example.com>",
        "To": "Bob <bob@example.com>",
        "Cc": "Carol <carol@example.com>",
        "Subject": "Quarterly numbers",
        "Date": "Tue, 1 Apr 2025 10:00:00 -0000",
    }

    fake_oxmsg = MagicMock()
    fake_oxmsg.Message.load.return_value = fake_msg

    loader = OutlookMsgLoader(str(file_path), include_headers=True)
    with patch.dict("sys.modules", {"oxmsg": fake_oxmsg}):
        docs = loader.load()

    assert len(docs) == 1
    text = docs[0].page_content
    assert "From: Alice <alice@example.com>" in text
    assert "To: Bob <bob@example.com>" in text
    assert "Cc: Carol <carol@example.com>" in text  # Cc preserved, not collapsed into To
    assert "Subject: Quarterly numbers" in text
    assert "Q1 results" in text
    fake_oxmsg.Message.load.assert_called_once_with(str(file_path))


def test_outlook_msg_loader_does_not_leak_bcc_recipients(tmp_path):
    """msg.recipients (which can include Bcc) must not be surfaced in the output."""
    from app.utils.document_loader import OutlookMsgLoader

    file_path = tmp_path / "mail.msg"
    file_path.write_bytes(b"placeholder")

    bcc = MagicMock(email_address="secret-bcc@example.com")
    bcc.name = "Secret"
    fake_msg = MagicMock()
    fake_msg.body = "Body."
    fake_msg.message_headers = {"To": "Bob <bob@example.com>"}
    fake_msg.recipients = [bcc]  # present, but must be ignored

    fake_oxmsg = MagicMock()
    fake_oxmsg.Message.load.return_value = fake_msg

    loader = OutlookMsgLoader(str(file_path), include_headers=True)
    with patch.dict("sys.modules", {"oxmsg": fake_oxmsg}):
        docs = loader.load()

    assert "secret-bcc@example.com" not in docs[0].page_content
    assert "To: Bob <bob@example.com>" in docs[0].page_content


def test_outlook_msg_loader_falls_back_to_attributes_without_headers(tmp_path):
    """When transport headers are absent, From/Subject/Date fall back to attributes."""
    from app.utils.document_loader import OutlookMsgLoader

    file_path = tmp_path / "mail.msg"
    file_path.write_bytes(b"placeholder")

    fake_msg = MagicMock()
    fake_msg.body = "Body text."
    fake_msg.message_headers = {}  # no transport headers (e.g. a sent draft)
    fake_msg.sender = "Alice <alice@example.com>"
    fake_msg.subject = "Quarterly numbers"
    fake_msg.sent_date = "Tue, 1 Apr 2025 10:00:00"

    fake_oxmsg = MagicMock()
    fake_oxmsg.Message.load.return_value = fake_msg

    loader = OutlookMsgLoader(str(file_path), include_headers=True)
    with patch.dict("sys.modules", {"oxmsg": fake_oxmsg}):
        docs = loader.load()

    text = docs[0].page_content
    assert "From: Alice <alice@example.com>" in text
    assert "Subject: Quarterly numbers" in text
    assert "Body text." in text


def test_outlook_msg_loader_body_only_when_headers_disabled(tmp_path):
    from app.utils.document_loader import OutlookMsgLoader

    file_path = tmp_path / "mail.msg"
    file_path.write_bytes(b"placeholder")

    fake_msg = MagicMock()
    fake_msg.body = "Just the body."

    fake_oxmsg = MagicMock()
    fake_oxmsg.Message.load.return_value = fake_msg

    loader = OutlookMsgLoader(str(file_path), include_headers=False)
    with patch.dict("sys.modules", {"oxmsg": fake_oxmsg}):
        docs = loader.load()

    assert docs[0].page_content == "Just the body."


# ---------------------------------------------------------------------------
# Standalone image OCR (.png/.jpg/.tiff/...)
# ---------------------------------------------------------------------------


def _make_png(path, color=(255, 0, 0)):
    from PIL import Image

    Image.new("RGB", (8, 8), color).save(path, format="PNG")


def _make_multiframe_tiff(path, frames=2):
    from PIL import Image

    imgs = [Image.new("RGB", (8, 8), (i * 10, 0, 0)) for i in range(frames)]
    imgs[0].save(path, format="TIFF", save_all=True, append_images=imgs[1:])


def test_get_loader_image_png(tmp_path):
    from app.utils.document_loader import ImageOCRLoader

    p = tmp_path / "scan.png"
    _make_png(str(p))

    loader, known_type, file_ext = get_loader("scan.png", "image/png", str(p))
    assert isinstance(loader, ImageOCRLoader)
    assert known_type is True
    assert file_ext == "png"


def test_get_loader_image_by_content_type(tmp_path):
    """Routing also works off an image/* Content-Type when the name lacks an ext."""
    from app.utils.document_loader import ImageOCRLoader

    p = tmp_path / "scan"
    _make_png(str(p))

    loader, _, _ = get_loader("scan", "image/tiff", str(p))
    assert isinstance(loader, ImageOCRLoader)


def test_get_loader_image_not_hijacked_by_markdown_mime(tmp_path):
    """A markdown Content-Type must not route an image into the markdown loader."""
    from app.utils.document_loader import ImageOCRLoader

    p = tmp_path / "scan.png"
    _make_png(str(p))

    loader, _, _ = get_loader("scan.png", "text/markdown", str(p))
    assert isinstance(loader, ImageOCRLoader)


def test_image_ocr_loader_single_image(tmp_path):
    from app.utils import document_loader
    from app.utils.document_loader import ImageOCRLoader

    p = tmp_path / "scan.png"
    _make_png(str(p))

    fake_module, client = _make_mistral_module(
        pages=[MagicMock(index=0, markdown="scanned text")]
    )

    loader = ImageOCRLoader(str(p))
    with patch.dict("sys.modules", {"mistralai": fake_module}), patch.object(
        document_loader, "MISTRAL_API_KEY", "test-key"
    ):
        docs = loader.load()

    assert len(docs) == 1
    assert docs[0].page_content == "scanned text"
    assert docs[0].metadata == {"source": str(p), "page": 1}
    # Sent to Mistral OCR as an image_url data URL
    _, kwargs = client.ocr.process.call_args
    assert kwargs["document"]["type"] == "image_url"
    assert kwargs["document"]["image_url"].startswith("data:image/png;base64,")


def test_image_ocr_loader_multipage_tiff_yields_one_doc_per_frame(tmp_path):
    from app.utils import document_loader
    from app.utils.document_loader import ImageOCRLoader

    p = tmp_path / "production.tiff"
    _make_multiframe_tiff(str(p), frames=2)

    fake_module, client = _make_mistral_module(
        pages=[MagicMock(index=0, markdown="page text")]
    )

    loader = ImageOCRLoader(str(p))
    with patch.dict("sys.modules", {"mistralai": fake_module}), patch.object(
        document_loader, "MISTRAL_API_KEY", "test-key"
    ):
        docs = loader.load()

    # One OCR call and one Document per TIFF frame, with sequential page numbers.
    assert client.ocr.process.call_count == 2
    assert [d.metadata["page"] for d in docs] == [1, 2]
    assert all(d.page_content == "page text" for d in docs)


def test_image_ocr_loader_caps_frames(tmp_path):
    """Frames beyond IMAGE_OCR_MAX_PAGES are skipped (no unbounded OCR fan-out)."""
    from app.utils import document_loader
    from app.utils.document_loader import ImageOCRLoader

    p = tmp_path / "huge.tiff"
    _make_multiframe_tiff(str(p), frames=5)

    fake_module, client = _make_mistral_module(
        pages=[MagicMock(index=0, markdown="page text")]
    )

    loader = ImageOCRLoader(str(p))
    with patch.dict("sys.modules", {"mistralai": fake_module}), patch.object(
        document_loader, "MISTRAL_API_KEY", "test-key"
    ), patch.object(document_loader, "IMAGE_OCR_MAX_PAGES", 2):
        docs = loader.load()

    assert client.ocr.process.call_count == 2
    assert len(docs) == 2


def test_image_ocr_loader_requires_api_key(tmp_path):
    from app.utils import document_loader
    from app.utils.document_loader import ImageOCRLoader

    p = tmp_path / "scan.png"
    _make_png(str(p))

    fake_module, _ = _make_mistral_module(pages=[])
    loader = ImageOCRLoader(str(p))
    with patch.dict("sys.modules", {"mistralai": fake_module}), patch.object(
        document_loader, "MISTRAL_API_KEY", ""
    ):
        with pytest.raises(RuntimeError, match="MISTRAL_API_KEY"):
            loader.load()
