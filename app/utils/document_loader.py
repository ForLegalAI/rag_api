# app/utils/document_loader.py

import os
import codecs
import tempfile

from typing import Iterator, List, Optional
import chardet

from langchain_core.documents import Document

from app.config import (
    known_source_ext,
    CHUNK_OVERLAP,
    logger,
    MISTRAL_API_KEY,
    MISTRAL_OCR_MODEL,
    DOCX_TEXT_USE_PANDOC,
    DOCX_TEXT_TRACK_CHANGES,
    DOCX_TEXT_INCLUDE_HEADERS_FOOTERS,
    EMAIL_INCLUDE_HEADERS,
    IMAGE_OCR_MAX_PAGES,
)
from langchain_community.document_loaders import (
    TextLoader,
    CSVLoader,
    Docx2txtLoader,
    UnstructuredEPubLoader,
    UnstructuredMarkdownLoader,
    UnstructuredXMLLoader,
    UnstructuredRSTLoader,
    UnstructuredRTFLoader,
    UnstructuredExcelLoader,
    UnstructuredPowerPointLoader,
    UnstructuredEmailLoader,
)


# Standalone raster image formats routed to OCR. SVG is excluded (vector/text,
# not raster-OCR friendly).
_IMAGE_EXTENSIONS = frozenset(
    {"png", "jpg", "jpeg", "gif", "bmp", "tif", "tiff", "webp"}
)

# Extensions that identify binary file formats handled by dedicated loaders.
# Used to prevent a conflicting multipart Content-Type (e.g. ``text/markdown``)
# from hijacking these files into a text loader. RTF is markup (not plain text),
# so it belongs here too, as do raster image formats.
_BINARY_FILE_EXTENSIONS = (
    frozenset(
        {"pdf", "doc", "docx", "xls", "xlsx", "ppt", "pptx", "epub", "msg", "eml", "rtf"}
    )
    | _IMAGE_EXTENSIONS
)


def detect_file_encoding(filepath: str) -> str:
    """
    Detect the encoding of a file using BOM markers and chardet for broader support.
    Returns the detected encoding or 'utf-8' as default.
    """
    with open(filepath, "rb") as f:
        raw = f.read(4096)  # Read a larger sample for better detection

    # Check for BOM markers first
    if raw.startswith(codecs.BOM_UTF16_LE):
        return "utf-16-le"
    elif raw.startswith(codecs.BOM_UTF16_BE):
        return "utf-16-be"
    elif raw.startswith(codecs.BOM_UTF16):
        return "utf-16"
    elif raw.startswith(codecs.BOM_UTF8):
        return "utf-8-sig"
    elif raw.startswith(codecs.BOM_UTF32_LE):
        return "utf-32-le"
    elif raw.startswith(codecs.BOM_UTF32_BE):
        return "utf-32-be"

    # Use chardet to detect encoding if no BOM is found
    result = chardet.detect(raw)
    encoding = result.get("encoding")
    if encoding:
        return encoding.lower()
    # Default to utf-8 if detection fails
    return "utf-8"


def cleanup_temp_encoding_file(loader) -> None:
    """
    Clean up temporary UTF-8 file if it was created for encoding conversion.

    :param loader: The document loader that may have created a temporary file
    """
    if hasattr(loader, "_temp_filepath") and loader._temp_filepath is not None:
        try:
            os.remove(loader._temp_filepath)
        except Exception as e:
            logger.warning(f"Failed to remove temporary UTF-8 file: {e}")


def get_loader(
    filename: str,
    file_content_type: str,
    filepath: str,
    raw_text: bool = False,
):
    """Get the appropriate document loader based on file type and/or content type.

    When ``raw_text`` is True, text-formatted files (e.g. Markdown) are loaded
    verbatim with :class:`TextLoader` so their original formatting is
    preserved. This is intended for the ``/text`` endpoint, where the caller
    wants the raw file contents. The embedding path should keep the default
    (``raw_text=False``) so semantic loaders continue to strip formatting for
    better vector search quality.
    """
    file_ext = filename.split(".")[-1].lower()
    known_type = True

    # File Content Type reference:
    # ref.: https://developer.mozilla.org/en-US/docs/Web/HTTP/Guides/MIME_types/Common_types
    if file_ext == "pdf" or file_content_type == "application/pdf":
        loader = SafePyPDFLoader(filepath)
    elif file_ext == "csv" or file_content_type == "text/csv":
        # Detect encoding for CSV files
        encoding = detect_file_encoding(filepath)

        if encoding != "utf-8":
            # For non-UTF-8 encodings, convert to UTF-8 using streaming
            # to avoid holding the entire file in memory as a single string
            temp_file = None
            try:
                with tempfile.NamedTemporaryFile(
                    mode="w", encoding="utf-8", suffix=".csv", delete=False
                ) as temp_file:
                    with open(
                        filepath, "r", encoding=encoding, errors="replace"
                    ) as original_file:
                        while True:
                            chunk = original_file.read(64 * 1024)
                            if not chunk:
                                break
                            temp_file.write(chunk)

                    temp_filepath = temp_file.name

                loader = CSVLoader(temp_filepath)
                loader._temp_filepath = temp_filepath
            except Exception as e:
                if temp_file and os.path.exists(temp_file.name):
                    os.unlink(temp_file.name)
                raise e
        else:
            loader = CSVLoader(filepath)
    elif file_ext == "rst":
        loader = UnstructuredRSTLoader(filepath, mode="elements")
    elif file_ext == "rtf" or file_content_type in ["application/rtf", "text/rtf"]:
        loader = UnstructuredRTFLoader(filepath)
    elif file_ext == "xml" or file_content_type in [
        "application/xml",
        "text/xml",
        "application/xhtml+xml",
    ]:
        loader = UnstructuredXMLLoader(filepath)
    elif file_ext in ["ppt", "pptx"] or file_content_type in [
        "application/vnd.ms-powerpoint",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ]:
        loader = UnstructuredPowerPointLoader(filepath)
    elif file_ext == "md" or (
        file_content_type
        in [
            "text/markdown",
            "text/x-markdown",
            "application/markdown",
            "application/x-markdown",
        ]
        and file_ext not in _BINARY_FILE_EXTENSIONS
    ):
        if raw_text:
            loader = TextLoader(filepath, autodetect_encoding=True)
        else:
            loader = UnstructuredMarkdownLoader(filepath)
    elif file_ext == "epub" or file_content_type == "application/epub+zip":
        loader = UnstructuredEPubLoader(filepath)
    elif file_ext == "doc" or file_content_type == "application/msword":
        # Legacy binary .doc (OLE2) is not supported: Docx2txtLoader only reads
        # OOXML .docx and pandoc can't read .doc either, so it could never load.
        # Reject clearly instead of failing later with a confusing error.
        raise ValueError(
            "Legacy .doc files are not supported. Please convert the document to .docx."
        )
    elif file_ext == "docx" or file_content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        # On the /text endpoint (raw_text=True), use pandoc so tracked changes,
        # comments and headers/footers are preserved; the embedding path uses
        # Docx2txtLoader.
        if raw_text and DOCX_TEXT_USE_PANDOC:
            loader = PandocDocxLoader(filepath)
        else:
            loader = Docx2txtLoader(filepath)
    elif file_ext == "eml" or file_content_type == "message/rfc822":
        loader = EmailLoader(filepath)
    elif file_ext == "msg" or file_content_type == "application/vnd.ms-outlook":
        loader = OutlookMsgLoader(filepath)
    elif file_ext in _IMAGE_EXTENSIONS or file_content_type in [
        "image/png",
        "image/jpeg",
        "image/gif",
        "image/bmp",
        "image/tiff",
        "image/webp",
    ]:
        # Standalone images (scanned exhibits, screenshots) run through OCR.
        loader = ImageOCRLoader(filepath)
    elif file_ext in ["xls", "xlsx"] or file_content_type in [
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ]:
        loader = UnstructuredExcelLoader(filepath)
    elif file_ext == "json" or file_content_type == "application/json":
        loader = TextLoader(filepath, autodetect_encoding=True)
    elif file_ext in known_source_ext or (
        file_content_type and file_content_type.find("text/") >= 0
    ):
        loader = TextLoader(filepath, autodetect_encoding=True)
    else:
        loader = TextLoader(filepath, autodetect_encoding=True)
        known_type = False

    return loader, known_type, file_ext


def clean_text(text: str) -> str:
    """
    Clean up text from PDF lopader

    :param text: The original text
    :return: Cleaned text
    """
    text = remove_null(text)
    text = remove_non_utf8(text)
    return text


def remove_null(text: str) -> str:
    """
    Remove NUL (0x00) characters from a string.

    :param text: The original text with potential NUL characters.
    :return: Cleaned text without NUL characters.
    """
    return text.replace("\x00", "")


def remove_non_utf8(text: str) -> str:
    """
    Remove invalid UTF-8 characters from a string, such as surrogate characters

    :param text: The original text with potential invalid utf-8 characters
    :return: Cleaned text without invalid utf-8 characters.
    """
    try:
        return text.encode("utf-8", "ignore").decode("utf-8")
    except UnicodeError:
        return text


def process_documents(documents: List[Document]) -> str:
    processed_text = ""
    last_page: Optional[int] = None
    doc_basename = ""

    for doc in documents:
        if "source" in doc.metadata:
            doc_basename = doc.metadata["source"].split("/")[-1]
            break

    processed_text += f"{doc_basename}\n"

    for doc in documents:
        current_page = doc.metadata.get("page")
        if current_page and current_page != last_page:
            processed_text += f"\n# PAGE {doc.metadata['page']}\n\n"
            last_page = current_page

        new_content = doc.page_content
        if processed_text.endswith(new_content[:CHUNK_OVERLAP]):
            processed_text += new_content[CHUNK_OVERLAP:]
        else:
            processed_text += new_content

    return processed_text.strip()


def _mistral_ocr_client():
    """Import mistralai, validate the API key, and return a configured client.

    Raised before any file encoding so OCR work is skipped when it can't run,
    and so a multi-page image builds the client only once.
    """
    # Lazy import to avoid hard dependency at import time
    try:
        from mistralai import Mistral
    except Exception as e:
        raise RuntimeError(
            "mistralai package is required for OCR. Please install 'mistralai' and set MISTRAL_API_KEY."
        ) from e

    if not MISTRAL_API_KEY:
        raise RuntimeError(
            "MISTRAL_API_KEY is not set. Please configure it in environment variables."
        )

    return Mistral(api_key=MISTRAL_API_KEY)


def _ocr_document(client, document_payload: dict, source: str, start_page: int = 1) -> List[Document]:
    """Run one Mistral OCR request and map its pages to Documents.

    Pages are numbered sequentially (1-based) from ``start_page`` in response
    order, so callers don't depend on the provider's own (0-based) page index.
    Returns a single empty placeholder when the response has no pages.
    """
    try:
        ocr_response = client.ocr.process(
            model=MISTRAL_OCR_MODEL,
            document=document_payload,
            include_image_base64=False,
        )
    except Exception as e:
        logger.error(f"Mistral OCR API call failed: {e}")
        raise

    pages = getattr(ocr_response, "pages", None)
    # Some clients return dict-like response; handle both
    if pages is None and isinstance(ocr_response, dict):
        pages = ocr_response.get("pages")

    if not pages:
        # Return an empty single document to avoid downstream crashes
        return [Document(page_content="", metadata={"source": source, "page": start_page})]

    documents: List[Document] = []
    for offset, page in enumerate(pages):
        markdown = (
            page.get("markdown") if isinstance(page, dict) else getattr(page, "markdown", None)
        )
        documents.append(
            Document(
                page_content=markdown or "",
                metadata={"source": source, "page": start_page + offset},
            )
        )

    return documents


class SafePyPDFLoader:
    """
    Replacement for previous PyPDF-based loader that now uses Mistral OCR API.
    Keeps the class name for compatibility with existing imports/tests.

    It returns one Document per page with metadata similar to PyPDFLoader:
    - metadata.source: original filepath
    - metadata.page: 1-based page index

    Images embedded in the PDF are not extracted.
    """

    def __init__(self, filepath: str):
        self.filepath = filepath
        self._temp_filepath = None  # For compatibility with cleanup function

    def _encode_pdf_b64(self) -> str:
        import base64

        with open(self.filepath, "rb") as f:
            return base64.b64encode(f.read()).decode("utf-8")

    def load(self) -> List[Document]:
        client = _mistral_ocr_client()  # validates the key before encoding the file
        payload = {
            "type": "document_url",
            "document_url": f"data:application/pdf;base64,{self._encode_pdf_b64()}",
        }
        return _ocr_document(client, payload, self.filepath)

    def lazy_load(self) -> Iterator[Document]:
        """Yield OCR-extracted pages one at a time.

        Mistral OCR returns the whole document in a single response, so true
        streaming isn't possible; this simply yields from :meth:`load`. It
        exists to satisfy the loader interface used by the document routes,
        which call ``lazy_load()``.
        """
        yield from self.load()


class ImageOCRLoader:
    """OCR a standalone image file via Mistral OCR.

    For images uploaded directly (scanned exhibits, screenshots, photos of
    documents) — not for images embedded inside PDFs. Single-frame PNG/JPEG are
    sent as-is; other formats and modes are normalized to PNG via Pillow, and
    multi-frame images (e.g. TIFF productions, animated GIF) yield one Document
    per frame, capped at ``IMAGE_OCR_MAX_PAGES``.
    """

    def __init__(self, filepath: str):
        self.filepath = filepath
        self._temp_filepath = None  # For compatibility with cleanup function

    def _iter_image_payloads(self) -> Iterator[tuple]:
        """Yield ``(mime, base64)`` for each frame to OCR.

        A single-frame PNG/JPEG is passed through unmodified (no decode/re-encode);
        everything else (multi-frame TIFF/GIF, palette/CMYK/RGBA modes) is
        normalized to PNG via Pillow.
        """
        import base64
        import io
        from PIL import Image, ImageSequence

        with Image.open(self.filepath) as img:
            fmt = (img.format or "").upper()
            if getattr(img, "n_frames", 1) == 1 and fmt in ("PNG", "JPEG"):
                with open(self.filepath, "rb") as f:
                    raw = f.read()
                mime = "image/png" if fmt == "PNG" else "image/jpeg"
                yield mime, base64.b64encode(raw).decode("utf-8")
                return

            for frame in ImageSequence.Iterator(img):
                # PNG supports RGB/L directly; convert other modes (P, RGBA,
                # CMYK, ...) to RGB so the encode never fails.
                normalized = frame if frame.mode in ("RGB", "L") else frame.convert("RGB")
                buf = io.BytesIO()
                normalized.save(buf, format="PNG")
                yield "image/png", base64.b64encode(buf.getvalue()).decode("utf-8")

    def load(self) -> List[Document]:
        client = _mistral_ocr_client()  # validates the key before decoding the image

        documents: List[Document] = []
        for mime, b64 in self._iter_image_payloads():
            if len(documents) >= IMAGE_OCR_MAX_PAGES:
                logger.warning(
                    "Image %s exceeds IMAGE_OCR_MAX_PAGES=%d; remaining frames skipped",
                    self.filepath,
                    IMAGE_OCR_MAX_PAGES,
                )
                break
            payload = {"type": "image_url", "image_url": f"data:{mime};base64,{b64}"}
            documents.extend(
                _ocr_document(client, payload, self.filepath, start_page=len(documents) + 1)
            )

        if not documents:
            # Preserve the always-at-least-one-Document invariant.
            return [Document(page_content="", metadata={"source": self.filepath, "page": 1})]
        return documents

    def lazy_load(self) -> Iterator[Document]:
        yield from self.load()


def _format_email_headers(headers: dict) -> str:
    """Build a header block from available email header values.

    Renders ``Label: value`` lines for the standard headers, skipping any that
    are missing/empty. Returns an empty string when no headers are present.
    """
    lines = []
    for label in ("From", "To", "Cc", "Subject", "Date"):
        value = headers.get(label)
        if value:
            lines.append(f"{label}: {value}")
    return "\n".join(lines)


def _combine_headers_and_body(header_block: str, body: str) -> str:
    """Join a header block and body with a blank line, tolerating empty sides."""
    if header_block and body:
        return f"{header_block}\n\n{body}"
    return header_block or body


def _extract_docx_headers_footers(filepath: str) -> str:
    """Extract distinct header/footer text from a .docx via python-docx.

    pandoc drops .docx headers/footers, but in legal documents they carry
    significant content (matter numbers, "PRIVILEGED & CONFIDENTIAL", "DRAFT").
    Returns a block of ``[Header] ...`` / ``[Footer] ...`` lines, deduplicated
    across sections and first/even-page variants. Returns "" on any failure so
    body extraction is never blocked.
    """
    try:
        from docx import Document as DocxDocument
    except Exception as e:
        logger.warning(f"python-docx unavailable; skipping header/footer extraction: {e}")
        return ""

    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        logger.warning(f"Failed to read .docx headers/footers for {filepath}: {e}")
        return ""

    lines = []
    seen = set()
    for section in doc.sections:
        parts = (
            (section.header, "Header"),
            (section.first_page_header, "Header"),
            (section.even_page_header, "Header"),
            (section.footer, "Footer"),
            (section.first_page_footer, "Footer"),
            (section.even_page_footer, "Footer"),
        )
        for hf, label in parts:
            text = "\n".join(p.text for p in hf.paragraphs).strip()
            if text and text not in seen:
                seen.add(text)
                lines.append(f"[{label}] {text}")
    return "\n".join(lines)


class PandocDocxLoader:
    """Load a ``.docx`` via pandoc so tracked changes and comments are preserved.

    Used by the ``/text`` endpoint (``raw_text=True``); the embedding path keeps
    using ``Docx2txtLoader``. Produces a single :class:`Document` whose
    ``page_content`` is pandoc Markdown. ``track_changes`` maps to pandoc's
    ``--track-changes`` flag (``all`` keeps insertions/deletions and, with the
    Markdown output used here, records each edit's author/date plus comments).
    pandoc only reads OOXML ``.docx``, not legacy binary ``.doc``.

    When ``include_headers_footers`` is enabled, header/footer text (which pandoc
    drops) is extracted separately via python-docx and prepended to the body.
    """

    def __init__(
        self,
        filepath: str,
        track_changes: str = DOCX_TEXT_TRACK_CHANGES,
        include_headers_footers: bool = DOCX_TEXT_INCLUDE_HEADERS_FOOTERS,
    ):
        self.filepath = filepath
        self.track_changes = track_changes
        self.include_headers_footers = include_headers_footers
        self._temp_filepath = None  # For compatibility with cleanup function

    def load(self) -> List[Document]:
        try:
            import pypandoc
        except Exception as e:
            raise RuntimeError(
                "pypandoc is required to extract .docx text via pandoc. "
                "Please install 'pypandoc' and ensure the pandoc binary is available."
            ) from e

        # Lets pypandoc's "No pandoc was found" OSError propagate so the /text
        # route can surface ERROR_MESSAGES.PANDOC_NOT_INSTALLED.
        text = pypandoc.convert_file(
            self.filepath,
            to="markdown",
            format="docx",
            extra_args=[f"--track-changes={self.track_changes}", "--wrap=none"],
        )

        if self.include_headers_footers:
            header_block = _extract_docx_headers_footers(self.filepath)
            text = _combine_headers_and_body(header_block, text)

        return [Document(page_content=text, metadata={"source": self.filepath})]

    def lazy_load(self) -> Iterator[Document]:
        yield from self.load()


class EmailLoader:
    """Load a ``.eml`` (RFC-822) message as a single Document.

    The body is extracted with :class:`UnstructuredEmailLoader` (handles plain
    and HTML parts); when ``include_headers`` is enabled, the standard headers
    (From/To/Cc/Subject/Date) parsed from the message are prepended.
    """

    def __init__(self, filepath: str, include_headers: bool = EMAIL_INCLUDE_HEADERS):
        self.filepath = filepath
        self.include_headers = include_headers
        self._temp_filepath = None  # For compatibility with cleanup function

    def _read_headers(self) -> str:
        from email import policy
        from email.parser import BytesParser

        # headersonly avoids re-parsing the body, which UnstructuredEmailLoader
        # already parsed for us.
        with open(self.filepath, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f, headersonly=True)
        return _format_email_headers(
            {
                "From": msg.get("From"),
                "To": msg.get("To"),
                "Cc": msg.get("Cc"),
                "Subject": msg.get("Subject"),
                "Date": msg.get("Date"),
            }
        )

    def load(self) -> List[Document]:
        body_docs = UnstructuredEmailLoader(self.filepath).load()
        body = "\n".join(d.page_content for d in body_docs).strip()

        content = body
        if self.include_headers:
            content = _combine_headers_and_body(self._read_headers(), body)

        return [Document(page_content=content, metadata={"source": self.filepath})]

    def lazy_load(self) -> Iterator[Document]:
        yield from self.load()


class OutlookMsgLoader:
    """Load a ``.msg`` (Outlook) message as a single Document via python-oxmsg.

    Extracts the plain-text body and, when ``include_headers`` is enabled,
    prepends From/To/Cc/Subject/Date.

    Header values come from the message's transport headers, which carry the
    real To/Cc split and — importantly — do **not** include Bcc. We deliberately
    avoid ``msg.recipients`` for this, because oxmsg exposes no recipient type,
    so using it would both collapse To/Cc and leak Bcc recipients (present in
    Sent-Items .msg files) into the extracted text. From/Subject/Date fall back
    to the structured attributes when a transport header is absent.
    """

    def __init__(self, filepath: str, include_headers: bool = EMAIL_INCLUDE_HEADERS):
        self.filepath = filepath
        self.include_headers = include_headers
        self._temp_filepath = None  # For compatibility with cleanup function

    def load(self) -> List[Document]:
        try:
            from oxmsg import Message
        except Exception as e:
            raise RuntimeError(
                "python-oxmsg is required to extract .msg files. "
                "Please install 'python-oxmsg'."
            ) from e

        msg = Message.load(self.filepath)
        body = (msg.body or "").strip()

        content = body
        if self.include_headers:
            raw_headers = getattr(msg, "message_headers", None) or {}
            headers = {str(k).lower(): v for k, v in raw_headers.items()}

            def hdr(name, fallback=None):
                return headers.get(name.lower()) or fallback

            sent_date = getattr(msg, "sent_date", None)
            header_block = _format_email_headers(
                {
                    "From": hdr("From", getattr(msg, "sender", None)),
                    "To": hdr("To"),
                    "Cc": hdr("Cc"),
                    "Subject": hdr("Subject", getattr(msg, "subject", None)),
                    "Date": hdr("Date", str(sent_date) if sent_date else None),
                }
            )
            content = _combine_headers_and_body(header_block, body)

        return [Document(page_content=content, metadata={"source": self.filepath})]

    def lazy_load(self) -> Iterator[Document]:
        yield from self.load()
